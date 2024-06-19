import re
import requests
from bs4 import BeautifulSoup
from docx import Document

# Function to extract email addresses from a webpage
def extract_emails(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # Check for HTTP errors
        soup = BeautifulSoup(response.text, 'html.parser')
        emails = set(re.findall(r'[a-zA-Z0-9._%+-]+@aol\.com', soup.text))
        return emails
    except requests.exceptions.RequestException as e:
        print(f"Error accessing {url}: {e}")
        return set()

# List of URLs to scrape (add more URLs as needed)
urls = [
    "https://branches.com.ng/contact",
    "http://example.com/page2",
    # Add more URLs here
]

# Set to store unique email addresses
email_addresses = set()

# Scrape each URL
for url in urls:
    print(f"Scraping {url}...")
    emails = extract_emails(url)
    email_addresses.update(emails)

# Create a new Word document
doc = Document()

# Add a title to the document
doc.add_heading('List of AOL Email Addresses', level=1)

# Add email addresses to the document with numbering
for i, email in enumerate(email_addresses, start=1):
    doc.add_paragraph(f"{i}. {email}")

# Save the document
doc.save('AOL_Email_Addresses.docx')

print("Document created successfully!")
