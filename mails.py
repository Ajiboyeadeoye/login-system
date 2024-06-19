import random
import string
from docx import Document

# Expanded lists of common first and last names to add more variability
first_names = [
    "john", "jane", "alex", "chris", "pat", "sam", "kelly", "morgan", "casey", "jordan",
    "taylor", "jamie", "drew", "ryan", "michele", "lee", "kim", "jessica", "michael", "emily",
    "daniel", "sarah", "matthew", "laura", "andrew", "megan", "david", "rachel", "joshua", "hannah",
    "nicholas", "ashley", "joseph", "brianna", "tyler", "olivia", "zachary", "madison", "justin", "emma"
]

last_names = [
    "smith", "johnson", "williams", "brown", "jones", "garcia", "miller", "davis", "rodriguez", "martinez",
    "hernandez", "lopez", "gonzalez", "wilson", "anderson", "thomas", "moore", "taylor", "lee", "white",
    "harris", "clark", "lewis", "robinson", "walker", "young", "allen", "king", "wright", "scott",
    "torres", "nguyen", "hill", "flores", "green", "adams", "nelson", "baker", "hall", "rivera"
]

company_names = [
    "techcorp", "bizsolutions", "innovatech", "webworld", "networks", "solarsystems", "greenenergy", "smartservices",
    "cloudcomputing", "datadynamics", "infotech", "cyberdyne", "softworks", "digitech", "techbridge", "alphatech",
    "nextgen", "synergy", "byteworks", "infosystems"
]

# Function to generate a random personal email address
def generate_personal_email():
    first_name = random.choice(first_names)
    last_name = random.choice(last_names)
    first_initial = first_name[0]
    last_initial = last_name[0]
    number = ''.join(random.choices(string.digits, k=random.randint(1, 4)))  # Adds a random length number between 1 and 4 digits
    email_format = random.choice([
        f"{first_name}{last_name}{number}",
        f"{first_name}.{last_name}{number}",
        f"{first_name}_{last_name}{number}",
        f"{first_name}{number}.{last_name}",
        f"{first_name}{number}_{last_name}",
        f"{first_initial}{last_name}{number}",
        f"{first_name}{last_initial}{number}",
        f"{first_initial}.{last_name}{number}",
        f"{first_name}.{last_initial}{number}"
    ])
    return f"{email_format}@aol.com"

# Function to generate a random company-like email address
def generate_company_email():
    first_name = random.choice(first_names)
    last_name = random.choice(last_names)
    first_initial = first_name[0]
    company = random.choice(company_names)
    number = ''.join(random.choices(string.digits, k=random.randint(1, 3)))  # Adds a random length number between 1 and 3 digits
    email_format = random.choice([
        f"{first_name}@{company}.aol.com",
        f"{first_initial}.{last_name}@{company}.aol.com",
        f"{first_name}.{number}@{company}.aol.com",
        f"{first_name}_{number}@{company}.aol.com"
    ])
    return email_format

# Create a new Word document
doc = Document()

# Add a title to the document
doc.add_heading('List of AOL Email Addresses', level=1)

# Generate 2000 email addresses and add them to the document with numbering
for i in range(1, 2001):
    if random.random() < 0.3:  # 30% chance to create a company-like email
        email = generate_company_email()
    else:
        email = generate_personal_email()
    doc.add_paragraph(f"{i}. {email}")

# Save the document
doc.save('AOL_Email_Addresses.docx')

print("Document created successfully!")
