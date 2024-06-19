
from docx import Document
import random
import string

# List of example names for more realistic emails
first_names = ["john", "jane", "alex", "emily", "michael", "sarah", "david", "laura", "chris", "megan"]
last_names = ["smith", "johnson", "williams", "brown", "jones", "miller", "davis", "garcia", "rodriguez", "wilson"]

def generate_realistic_email():
    first_name = random.choice(first_names)
    last_name = random.choice(last_names)
    number = random.randint(1, 99)  # Adds a number to the email for more realism
    domain = "aol.com"
    return f"{first_name}.{last_name}{number}@{domain}"

# Create a Word document
doc = Document()
doc.add_heading('List of Realistic Fictional AOL Email Addresses', 0)

# Generate 7000 email addresses and add them to the document
for _ in range(7000):
    email = generate_realistic_email()
    doc.add_paragraph(email)

# Save the document
doc.save('realistic_fictional_aol_emails.docx')

print("2000 realistic fictional AOL email addresses have been generated and saved to realistic_fictional_aol_emails.docx")
